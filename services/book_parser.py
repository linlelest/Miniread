"""
Miniread (极读) - 电子书解析服务
支持多种格式的解析：EPUB, FB2, DOCX, HTML, MD, RTF, TXT, MOBI等
"""
import os
import re
import zipfile
import tempfile
import shutil
from config import Config


def parse_epub(file_path):
    """
    解析EPUB文件
    返回: {title, author, chapters, epub_chapters, total_chapters}
    """
    result = {
        'title': '',
        'author': '',
        'chapters': [],
        'epub_chapters': [],
        'total_chapters': 0,
    }

    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(file_path)

        # 提取元数据
        title_list = book.get_metadata('DC', 'title')
        if title_list:
            result['title'] = title_list[0][0]
        else:
            result['title'] = os.path.splitext(os.path.basename(file_path))[0]

        creator_list = book.get_metadata('DC', 'creator')
        if creator_list:
            result['author'] = creator_list[0][0]

        # 提取章节 — spine entries are (item_id, show) tuples
        chapters = []
        epub_chapters = []
        spine = book.spine if hasattr(book, 'spine') else []

        # Build spine-order index for sorting items
        spine_keys = {}
        for idx, entry in enumerate(spine):
            if isinstance(entry, tuple) and len(entry) >= 1:
                spine_keys[entry[0]] = idx

        # Get all items sorted by spine position
        all_items = list(book.get_items())
        all_items.sort(key=lambda it: spine_keys.get(it.get_id(), 9999))

        for i, item in enumerate(all_items):
            if item.get_type() != ebooklib.ITEM_DOCUMENT:
                continue
            try:
                content = item.get_content()
                if isinstance(content, bytes):
                    try:
                        content = content.decode('utf-8')
                    except UnicodeDecodeError:
                        content = content.decode('latin-1', errors='replace')
            except:
                continue

            soup = BeautifulSoup(content, 'html.parser')

            # 尝试提取标题
            title_tag = soup.find(['h1', 'h2', 'h3', 'h4'])
            if title_tag:
                chapter_title = title_tag.get_text(strip=True)
            else:
                chapter_title = f'章节 {i + 1}'

            # 提取正文
            body = soup.find('body')
            if body:
                for tag in body.find_all(['script', 'style']):
                    tag.decompose()
                chapter_html = str(body)
            else:
                chapter_html = str(soup)

            chapters.append({
                'title': chapter_title[:100],
                'index': i,
                'position': i / max(len(spine), 1),
            })

            epub_chapters.append({
                'title': chapter_title[:100],
                'content': chapter_html,
                'index': i,
            })

        result['chapters'] = chapters
        result['epub_chapters'] = epub_chapters
        result['total_chapters'] = len(chapters)

    except ImportError:
        # ebooklib不可用时的回退方案
        return _parse_epub_fallback(file_path)
    except Exception as e:
        # 解析失败，返回基本信息
        result['title'] = os.path.splitext(os.path.basename(file_path))[0]
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'], 'content': f'<p>EPUB解析错误: {str(e)}</p>', 'index': 0}]
        result['total_chapters'] = 1

    return result


def _parse_epub_fallback(file_path):
    """EPUB回退解析方案（基本zip + XML解析）"""
    result = {
        'title': os.path.splitext(os.path.basename(file_path))[0],
        'author': '',
        'chapters': [],
        'epub_chapters': [],
        'total_chapters': 0,
    }

    try:
        with zipfile.ZipFile(file_path, 'r') as zf:
            # 查找OPF文件
            opf_path = None
            for name in zf.namelist():
                if name.endswith('.opf'):
                    opf_path = name
                    break

            if opf_path:
                opf_content = zf.read(opf_path).decode('utf-8', errors='replace')
                # 提取标题
                title_match = re.search(r'<dc:title[^>]*>([^<]+)</dc:title>', opf_content, re.IGNORECASE)
                if title_match:
                    result['title'] = title_match.group(1)
                # 提取作者
                author_match = re.search(r'<dc:creator[^>]*>([^<]+)</dc:creator>', opf_content, re.IGNORECASE)
                if author_match:
                    result['author'] = author_match.group(1)

            # 提取HTML/XHTML内容
            html_files = [n for n in zf.namelist()
                          if n.lower().endswith(('.html', '.htm', '.xhtml'))]
            html_files.sort()

            from bs4 import BeautifulSoup
            chapters = []
            epub_chapters = []

            for i, html_file in enumerate(html_files):
                try:
                    content = zf.read(html_file).decode('utf-8', errors='replace')
                except:
                    content = zf.read(html_file).decode('latin-1', errors='replace')

                soup = BeautifulSoup(content, 'html.parser')

                title_tag = soup.find(['h1', 'h2', 'h3', 'h4', 'title'])
                chapter_title = title_tag.get_text(strip=True)[:100] if title_tag else f'章节 {i + 1}'

                body = soup.find('body')
                chapter_html = str(body) if body else str(soup)

                chapters.append({
                    'title': chapter_title,
                    'index': i,
                    'position': i / max(len(html_files), 1),
                })
                epub_chapters.append({
                    'title': chapter_title,
                    'content': chapter_html,
                    'index': i,
                })

        # Fallback if no chapters detected from spine
        if not chapters:
            from utils.helpers import detect_chapters_txt
            # Try to find any text content in the book
            all_text = ''
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    try:
                        all_text += item.get_content().decode('utf-8', errors='replace')
                    except:
                        pass
            if all_text:
                lines = [l for l in all_text.split('\n') if l.strip()]
                chunk_size = max(1, len(lines) // 20)
                for i in range(0, len(lines), chunk_size):
                    end = min(i + chunk_size, len(lines))
                    chunk_title = lines[i].strip()[:50] or f'第{len(chapters)+1}节'
                    chapters.append({
                        'title': chunk_title,
                        'index': len(chapters),
                        'position': i / max(len(lines), 1)
                    })
                    epub_chapters.append({
                        'title': chunk_title,
                        'content': '<div>' + '</p><p>'.join(
                            l.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
                            for l in lines[i:end] if l.strip()
                        ) + '</div>',
                        'index': len(epub_chapters)
                    })

        result['chapters'] = chapters
        result['epub_chapters'] = epub_chapters
        result['total_chapters'] = len(chapters)

    except Exception as e:
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'], 'content': f'<p>文件读取失败: {str(e)}</p>', 'index': 0}]
        result['total_chapters'] = 1

    if not result.get('chapters'):
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'], 'content': '<p>章节提取失败</p>', 'index': 0}]
        result['total_chapters'] = 1

    return result


def extract_epub_cover(file_path):
    """从EPUB中提取封面图片 — 多种策略"""
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
        import base64

        book = epub.read_epub(file_path)
        cover_dir = os.path.join(tempfile.gettempdir(), 'miniread_covers')
        os.makedirs(cover_dir, exist_ok=True)

        # Strategy 1: Check by ITEM_COVER type or name containing 'cover'
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_COVER or \
               (item.get_type() == ebooklib.ITEM_IMAGE and 'cover' in item.get_name().lower()):
                ext = os.path.splitext(item.get_name())[1].lower()
                if ext in ('.jpg', '.jpeg', '.png', '.gif'):
                    cover_path = os.path.join(cover_dir, f'cover_{os.path.basename(file_path)}{ext}')
                    with open(cover_path, 'wb') as f:
                        f.write(item.get_content())
                    return cover_path

        # Strategy 2: First image in spine order
        spine_order = {}
        if hasattr(book, 'spine'):
            for idx, entry in enumerate(book.spine):
                if isinstance(entry, tuple) and len(entry) >= 1:
                    spine_order[entry[0]] = idx

        all_items = list(book.get_items())
        all_items.sort(key=lambda it: spine_order.get(it.get_id(), 9999))

        # Find documents first, then check for images in the first doc
        for item in all_items:
            if item.get_type() == ebooklib.ITEM_IMAGE:
                ext = os.path.splitext(item.get_name())[1].lower()
                if ext in ('.jpg', '.jpeg', '.png', '.gif', '.svg'):
                    cover_path = os.path.join(cover_dir, f'cover_{os.path.basename(file_path)}{ext}')
                    with open(cover_path, 'wb') as f:
                        f.write(item.get_content())
                    return cover_path

            # Check first document for img tags
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                try:
                    raw = item.get_content()
                    if isinstance(raw, bytes):
                        text = raw.decode('utf-8', errors='replace')
                    else:
                        text = raw
                    soup = BeautifulSoup(text, 'html.parser')
                    img = soup.find('img')
                    if img and img.get('src'):
                        src = img['src'].split('/')[-1]
                        # Find this image in the book
                        for img_item in book.get_items():
                            if img_item.get_type() == ebooklib.ITEM_IMAGE and \
                               (img_item.get_name().endswith(src) or os.path.basename(img_item.get_name()) == src):
                                ext = os.path.splitext(img_item.get_name())[1].lower()
                                cover_path = os.path.join(cover_dir, f'cover_{os.path.basename(file_path)}{ext}')
                                with open(cover_path, 'wb') as f:
                                    f.write(img_item.get_content())
                                return cover_path
                except:
                    pass
                break  # Only check first doc
    except:
        pass
    return None


def parse_fb2(file_path):
    """
    解析FB2 (FictionBook) 文件
    FB2是XML格式
    """
    result = {
        'title': os.path.splitext(os.path.basename(file_path))[0],
        'author': '',
        'chapters': [],
        'epub_chapters': [],
        'total_chapters': 0,
    }

    try:
        from bs4 import BeautifulSoup

        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        soup = BeautifulSoup(content, 'xml')
        # 或使用 html.parser
        if not soup.find('FictionBook'):
            soup = BeautifulSoup(content, 'html.parser')

        # 提取标题
        book_title = soup.find('book-title')
        if book_title:
            result['title'] = book_title.get_text(strip=True)

        # 提取作者
        first_name = soup.find('first-name')
        last_name = soup.find('last-name')
        if first_name or last_name:
            result['author'] = f"{first_name.get_text(strip=True) if first_name else ''} {last_name.get_text(strip=True) if last_name else ''}".strip()

        # 提取章节 (section)
        body = soup.find('body')
        sections = body.find_all('section') if body else []

        chapters = []
        epub_chapters = []

        if not sections:
            # 整个文档作为一个章节
            text_content = str(body) if body else str(soup)
            chapters.append({
                'title': result['title'],
                'index': 0,
                'position': 0,
            })
            epub_chapters.append({
                'title': result['title'],
                'content': _fb2_to_html(text_content),
                'index': 0,
            })
        else:
            for i, section in enumerate(sections):
                title_tag = section.find('title')
                chapter_title = title_tag.get_text(strip=True)[:100] if title_tag else f'章节 {i + 1}'

                # 转换为HTML
                chapter_html = _fb2_to_html(str(section))

                chapters.append({
                    'title': chapter_title,
                    'index': i,
                    'position': i / max(len(sections), 1),
                })
                epub_chapters.append({
                    'title': chapter_title,
                    'content': chapter_html,
                    'index': i,
                })

        result['chapters'] = chapters
        result['epub_chapters'] = epub_chapters
        result['total_chapters'] = len(chapters)

    except Exception as e:
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'], 'content': f'<p>FB2解析失败: {str(e)}</p>', 'index': 0}]
        result['total_chapters'] = 1

    return result


def _fb2_to_html(fb2_xml):
    """将FB2 XML内容转换为HTML"""
    html = fb2_xml
    # 转换常见的FB2标签
    replacements = [
        ('<section>', '<div>'), ('</section>', '</div>'),
        ('<title>', '<h3>'), ('</title>', '</h3>'),
        ('<subtitle>', '<h4>'), ('</subtitle>', '</h4>'),
        ('<p>', '<p>'), ('</p>', '</p>'),
        ('<emphasis>', '<i>'), ('</emphasis>', '</i>'),
        ('<strong>', '<b>'), ('</strong>', '</b>'),
        ('<strikethrough>', '<s>'), ('</strikethrough>', '</s>'),
        ('<epigraph>', '<blockquote>'), ('</epigraph>', '</blockquote>'),
        ('<cite>', '<cite>'), ('</cite>', '</cite>'),
        ('<poem>', '<div class="poem">'), ('</poem>', '</div>'),
        ('<stanza>', '<div class="stanza">'), ('</stanza>', '</div>'),
        ('<v>', '<p class="verse">'), ('</v>', '</p>'),
        ('<empty-line/>', '<br/>'),
        ('<image', '<img'),  # 简化处理
    ]
    for old, new in replacements:
        html = html.replace(old, new)

    return html


def parse_docx(file_path):
    """
    解析DOCX文件 — 使用 mammoth 完整保留Word样式
    支持：标题/H1-H6/列表/表格/图片/粗斜体/下划线/链接/对齐/缩进/行距
    """
    result = {
        'title': os.path.splitext(os.path.basename(file_path))[0],
        'author': '',
        'chapters': [],
        'epub_chapters': [],
        'total_chapters': 0,
    }

    try:
        import mammoth

        # Generate HTML with formatting
        style_map = """
        p[style-name='Title'] => h1:fresh
        p[style-name='Subtitle'] => h2:fresh
        p[style-name='Heading 1'] => h1:fresh
        p[style-name='Heading 2'] => h2:fresh
        p[style-name='Heading 3'] => h3:fresh
        p[style-name='Heading 4'] => h4:fresh
        p[style-name='heading 1'] => h1:fresh
        p[style-name='heading 2'] => h2:fresh
        p[style-name='heading 3'] => h3:fresh
        r[style-name='Strong'] => strong
        r[style-name='Emphasis'] => emph
        """

        with open(file_path, 'rb') as f:
            conv_result = mammoth.convert_to_html(f, style_map=style_map)
            full_html = conv_result.value
            # messages = conv_result.messages

        if not full_html.strip():
            raise ValueError('mammoth returned empty HTML')

        # Add styles for tables, alignment etc
        styled_html = '<div class="docx-content">' + full_html + '</div>'

        # Extract text for chapter detection
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(full_html, 'html.parser')
        all_text = soup.get_text(separator='\n')

        # Try to detect chapters
        from utils.helpers import detect_chapters_txt
        chapters_data = detect_chapters_txt(all_text)

        chapters = []
        epub_chapters = []

        if chapters_data and len(chapters_data) > 1:
            for i, ch in enumerate(chapters_data):
                ch_title = ch['title'][:100]
                chapters.append({
                    'title': ch_title, 'index': i,
                    'position': i / max(len(chapters_data), 1)
                })
                # For chapter content we serve the full rich HTML (mammoth handles all formatting)
                # Since mammoth produces clean HTML, serve it as one chapter
                epub_chapters.append({
                    'title': ch_title,
                    'content': styled_html,
                    'index': i
                })
        else:
            chapters = [{'title': result['title'], 'index': 0, 'position': 0}]
            epub_chapters = [{
                'title': result['title'],
                'content': styled_html,
                'index': 0
            }]

        result['chapters'] = chapters
        result['epub_chapters'] = epub_chapters
        result['total_chapters'] = len(chapters) or 1

    except ImportError:
        # mammoth not installed, try python-docx fallback
        return _parse_docx_rich_fallback(file_path)
    except Exception as e:
        # Fallback to python-docx
        try:
            return _parse_docx_rich_fallback(file_path)
        except:
            result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
            result['epub_chapters'] = [{'title': result['title'],
                'content': f'<p>DOCX解析失败: {esc(str(e))}</p>', 'index': 0}]
            result['total_chapters'] = 1

    if not result.get('chapters'):
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'],
            'content': '<p>章节提取失败</p>', 'index': 0}]
        result['total_chapters'] = 1

    return result


def _parse_docx_rich_fallback(file_path):
    """
    DOCX回退解析 — python-docx 全格式渲染
    保留：表格/图片/粗/斜/下划线/标题/列表
    """
    result = {
        'title': os.path.splitext(os.path.basename(file_path))[0],
        'author': '',
        'chapters': [],
        'epub_chapters': [],
        'total_chapters': 0,
    }

    try:
        from docx import Document
        from utils.helpers import detect_chapters_txt
        import base64, io

        doc = Document(file_path)

        def _run_html(run):
            t = esc(run.text)
            if run.bold: t = f'<b>{t}</b>'
            if run.italic: t = f'<i>{t}</i>'
            if run.underline: t = f'<u>{t}</u>'
            if run.font.strike: t = f'<s>{t}</s>'
            if run.font.superscript: t = f'<sup>{t}</sup>'
            if run.font.subscript: t = f'<sub>{t}</sub>'
            # Color
            if run.font.color and run.font.color.rgb:
                t = f'<span style="color:#{run.font.color.rgb}">{t}</span>'
            # Font size
            if run.font.size:
                try: t = f'<span style="font-size:{run.font.size.pt}pt">{t}</span>'
                except: pass
            return t

        def _para_html(para):
            if not para.text.strip():
                return ''
            style = para.style.name or ''
            sl = style.lower()
            # Headings
            if 'heading' in sl or 'title' in sl:
                level = 'h1' if '1' in sl else 'h2' if '2' in sl else 'h3' if '3' in sl else 'h4'
                inner = ''.join(_run_html(r) for r in para.runs)
                return f'<{level} style="margin:14px 0 6px;font-weight:600">'+inner+f'</{level}>'

            inner = ''.join(_run_html(r) for r in para.runs)
            align = ''
            if para.alignment:
                align_map = {0:'left',1:'center',2:'right',3:'both',4:'left'}
                a = align_map.get(para.alignment, '')
                if a: align = f' text-align:{a};'

            indent = ''
            if para.paragraph_format.left_indent:
                try: indent = f' padding-left:{para.paragraph_format.left_indent.pt}pt;'
                except: pass

            return f'<p style="margin:3px 0;line-height:1.7{align}{indent}">{inner}</p>'

        def _table_html(table):
            h = '<div class="docx-table-wrap" style="overflow-x:auto;margin:8px 0"><table style="width:100%;border-collapse:collapse;border:1px solid var(--border)">'
            for row in table.rows:
                h += '<tr>'
                for cell in row.cells:
                    # Handle merged cells
                    colspan = 1
                    tc = cell._tc
                    if tc.find('.{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan'):
                        import lxml.etree as et
                        gs = tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
                        if gs is not None:
                            colspan = int(gs.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '1'))
                    h += f'<td colspan="{colspan}" style="border:1px solid var(--border);padding:6px 10px;font-size:13px;vertical-align:top">{esc(cell.text.strip()[:200])}</td>'
                h += '</tr>'
            h += '</table></div>'
            return h

        # Extract images
        def _get_images():
            imgs = {}
            for rel in doc.part.rels.values():
                if 'image' in rel.reltype:
                    try:
                        b64 = base64.b64encode(rel.target_part.blob).decode()
                        ext = os.path.splitext(rel.target_ref)[1].lstrip('.') if hasattr(rel,'target_ref') else 'png'
                        if ext in ('jpg','jpeg'): ext='jpeg'
                        if ext in ('png','jpeg','gif','svg'):
                            imgs[rel.target_part.rId] = f'data:image/{ext};base64,{b64}'
                    except: pass
            return imgs

        images = _get_images()

        # Walk elements in order
        parts = ['<div>']
        texts = []
        for elem in doc.element.body:
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            if tag == 'p':
                for p in doc.paragraphs:
                    if p._element is elem:
                        h = _para_html(p)
                        if h: parts.append(h)
                        if p.text.strip(): texts.append(p.text)
                        break
            elif tag == 'tbl':
                for t in doc.tables:
                    parts.append(_table_html(t))
                    break

        parts.append('</div>')
        full_html = '\n'.join(parts)
        all_text = '\n'.join(texts)

        ch_data = detect_chapters_txt(all_text)
        if ch_data and len(ch_data) > 1:
            for i, ch in enumerate(ch_data):
                result['chapters'].append({'title': ch['title'][:100], 'index': i,
                    'position': i / max(len(ch_data), 1)})
                result['epub_chapters'].append({'title': ch['title'][:100],
                    'content': full_html, 'index': i})
        else:
            result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
            result['epub_chapters'] = [{'title': result['title'], 'content': full_html, 'index': 0}]

        result['total_chapters'] = len(result['chapters']) or 1

    except Exception as e:
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'],
            'content': f'<p>DOCX回退解析失败: {esc(str(e))}</p>', 'index': 0}]
        result['total_chapters'] = 1

    if not result.get('chapters'):
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'],
            'content': '<p>章节提取失败</p>', 'index': 0}]
        result['total_chapters'] = 1

    return result


def parse_html(file_path):
    """
    解析HTML文件
    """
    result = {
        'title': os.path.splitext(os.path.basename(file_path))[0],
        'author': '',
        'chapters': [],
        'epub_chapters': [],
        'total_chapters': 0,
    }

    try:
        from bs4 import BeautifulSoup

        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        soup = BeautifulSoup(content, 'html.parser')

        # 提取标题
        title_tag = soup.find('title')
        if title_tag:
            result['title'] = title_tag.get_text(strip=True)[:200]
        else:
            h1 = soup.find('h1')
            if h1:
                result['title'] = h1.get_text(strip=True)[:200]

        # 提取作者meta
        author_meta = soup.find('meta', {'name': 'author'})
        if author_meta:
            result['author'] = author_meta.get('content', '')

        # 查找章节标记 (h1-h4标签作为章节分隔)
        headers = soup.find_all(['h1', 'h2', 'h3', 'h4'])
        body = soup.find('body') or soup

        chapters = []
        epub_chapters = []

        if not headers:
            # 整个HTML作为一个章节
            for tag in body.find_all(['script', 'style', 'nav']):
                tag.decompose()
            chapters = [{'title': result['title'], 'index': 0, 'position': 0}]
            epub_chapters = [{'title': result['title'], 'content': str(body), 'index': 0}]
        else:
            # 按header分割
            for i, header in enumerate(headers):
                chapter_title = header.get_text(strip=True)[:100]
                # 收集该header之后到下一个header之间的内容
                content_parts = [str(header)]
                current = header.next_sibling
                while current and (i + 1 >= len(headers) or current != headers[i + 1]):
                    if hasattr(current, 'name') and current.name not in ('script', 'style'):
                        content_parts.append(str(current))
                    elif isinstance(current, str) and current.strip():
                        content_parts.append(f'<p>{current.strip()}</p>')
                    current = current.next_sibling

                chapters.append({
                    'title': chapter_title,
                    'index': i,
                    'position': i / max(len(headers), 1),
                })
                epub_chapters.append({
                    'title': chapter_title,
                    'content': '\n'.join(content_parts),
                    'index': i,
                })

        result['chapters'] = chapters
        result['epub_chapters'] = epub_chapters
        result['total_chapters'] = len(chapters)

    except Exception as e:
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'], 'content': f'<p>HTML解析失败: {str(e)}</p>', 'index': 0}]
        result['total_chapters'] = 1

    return result


def parse_markdown(file_path):
    """
    解析Markdown文件
    """
    result = {
        'title': os.path.splitext(os.path.basename(file_path))[0],
        'author': '',
        'chapters': [],
        'epub_chapters': [],
        'total_chapters': 0,
    }

    try:
        import markdown as md_lib

        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 使用markdown库转换整个文档
        html_content = md_lib.markdown(content, extensions=['extra', 'codehilite', 'toc'])

        # 按 ## 或 # 标题分割章节
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        headers = soup.find_all(['h1', 'h2', 'h3'])

        chapters = []
        epub_chapters = []

        if not headers:
            chapters = [{'title': result['title'], 'index': 0, 'position': 0}]
            epub_chapters = [{'title': result['title'], 'content': html_content, 'index': 0}]
        else:
            # 找到第一个header，设置为title
            result['title'] = headers[0].get_text(strip=True)[:200]

            for i, header in enumerate(headers):
                chapter_title = header.get_text(strip=True)[:100]
                # 使用BeautifulSoup重新解析以提取该章节
                parts = [str(header)]
                current = header.next_sibling
                while current and (i + 1 >= len(headers) or current != headers[i + 1]):
                    if hasattr(current, 'name'):
                        parts.append(str(current))
                    elif isinstance(current, str) and current.strip():
                        parts.append(f'<p>{current.strip()}</p>')
                    current = current.next_sibling

                chapters.append({
                    'title': chapter_title,
                    'index': i,
                    'position': i / max(len(headers), 1),
                })
                epub_chapters.append({
                    'title': chapter_title,
                    'content': '\n'.join(parts),
                    'index': i,
                })

        result['chapters'] = chapters
        result['epub_chapters'] = epub_chapters
        result['total_chapters'] = len(chapters)

    except ImportError:
        # markdown库不可用，作为纯文本处理
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'], 'content': f'<pre>{content}</pre>', 'index': 0}]
        result['total_chapters'] = 1
    except Exception as e:
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'], 'content': f'<p>Markdown解析失败: {str(e)}</p>', 'index': 0}]
        result['total_chapters'] = 1

    return result


def parse_rtf(file_path):
    """
    解析RTF文件
    """
    result = {
        'title': os.path.splitext(os.path.basename(file_path))[0],
        'author': '',
        'chapters': [],
        'epub_chapters': [],
        'total_chapters': 0,
    }

    try:
        from striprtf.striprtf import rtf_to_text

        with open(file_path, 'r', encoding='utf-8') as f:
            rtf_content = f.read()

        text = rtf_to_text(rtf_content)

        from utils.helpers import detect_chapters_txt
        chapters_data = detect_chapters_txt(text)

        chapters = []
        epub_chapters = []

        if chapters_data:
            for i, ch in enumerate(chapters_data):
                chapters.append({
                    'title': ch['title'][:100],
                    'index': i,
                    'position': i / max(len(chapters_data), 1),
                })
                start = ch['position']
                end = chapters_data[i + 1]['position'] if i + 1 < len(chapters_data) else len(text)
                seg_html = '<div>' + '</p><p>'.join(
                    line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    for line in text[start:end].split('\n') if line.strip()
                ) + '</div>'
                epub_chapters.append({
                    'title': ch['title'][:100],
                    'content': seg_html,
                    'index': i,
                })
        else:
            html = '<div>' + '</p><p>'.join(
                line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                for line in text.split('\n') if line.strip()
            ) + '</div>'
            chapters = [{'title': result['title'], 'index': 0, 'position': 0}]
            epub_chapters = [{'title': result['title'], 'content': html, 'index': 0}]

        result['chapters'] = chapters
        result['epub_chapters'] = epub_chapters
        result['total_chapters'] = len(chapters)

    except ImportError:
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'], 'content': '<p>RTF解析需要striprtf库</p>', 'index': 0}]
        result['total_chapters'] = 1
    except Exception as e:
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'], 'content': f'<p>RTF解析失败: {str(e)}</p>', 'index': 0}]
        result['total_chapters'] = 1

    return result


def parse_pdf(file_path):
    """
    解析PDF文件 — 使用pdfplumber提取文本
    """
    result = {
        'title': os.path.splitext(os.path.basename(file_path))[0],
        'author': '',
        'chapters': [],
        'epub_chapters': [],
        'total_chapters': 0,
    }

    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            all_pages_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    all_pages_text.append(text)

            full_text = '\n'.join(all_pages_text)

            if not full_text.strip():
                raise ValueError('PDF中未提取到文本（可能是扫描件）')

            # 尝试检测章节
            from utils.helpers import detect_chapters_txt
            chapters_data = detect_chapters_txt(full_text)

            chapters = []
            epub_chapters = []

            if chapters_data and len(chapters_data) > 1:
                for i, ch in enumerate(chapters_data):
                    chapters.append({
                        'title': ch['title'][:100],
                        'index': i,
                        'position': i / max(len(chapters_data), 1),
                    })
                    start = ch['position']
                    end = chapters_data[i + 1]['position'] if i + 1 < len(chapters_data) else len(full_text)
                    seg = full_text[start:end].strip()
                    seg_html = '<div>' + '</p><p>'.join(
                        esc(line) for line in seg.split('\n') if line.strip()
                    ) + '</div>'
                    epub_chapters.append({
                        'title': ch['title'][:100],
                        'content': seg_html,
                        'index': i,
                    })
            else:
                # 按页分章节
                for i, page_text in enumerate(all_pages_text):
                    if not page_text.strip():
                        continue
                    chapters.append({
                        'title': f'第{i + 1}页',
                        'index': i,
                        'position': i / max(len(all_pages_text), 1),
                    })
                    page_html = '<div>' + '</p><p>'.join(
                        esc(line) for line in page_text.split('\n') if line.strip()
                    ) + '</div>'
                    epub_chapters.append({
                        'title': f'第{i + 1}页',
                        'content': page_html,
                        'index': i,
                    })

            result['chapters'] = chapters or [{'title': result['title'], 'index': 0, 'position': 0}]
            result['epub_chapters'] = epub_chapters or [{'title': result['title'], 'content': '<p>无内容</p>', 'index': 0}]
            result['total_chapters'] = len(chapters) or 1

    except ImportError:
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'], 'content': '<p>PDF解析需要安装pdfplumber库: pip install pdfplumber</p>', 'index': 0}]
        result['total_chapters'] = 1
    except Exception as e:
        result['chapters'] = [{'title': result['title'], 'index': 0, 'position': 0}]
        result['epub_chapters'] = [{'title': result['title'], 'content': '<p>PDF解析失败: ' + esc(str(e)) + '</p>', 'index': 0}]
        result['total_chapters'] = 1

    return result


def esc(text):
    """HTML转义"""
    return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')


def extract_txt_preview(file_path, max_chars=500):
    """提取TXT文件预览"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read(max_chars)
        return text
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='gbk') as f:
            text = f.read(max_chars)
        return text
    except:
        return ''
