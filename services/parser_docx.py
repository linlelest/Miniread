import re
from html import escape


def _style_map():
    parts = []
    for n in range(1, 7):
        parts.append("p[style-name='标题 %d'] => h%d:fresh" % (n, n))
        parts.append("p[style-name='标题%d'] => h%d:fresh" % (n, n))
        parts.append("p[style-name='Heading %d'] => h%d:fresh" % (n, n))
    return '\n'.join(parts)


def _heading_level_of(style):
    if style is None:
        return None
    try:
        name = style.name or ''
        sid = style.style_id or ''
    except Exception:
        return None
    for text in (name, sid):
        m = re.search(r'(?:heading|标题)\s*(\d+)', text, re.I)
        if m:
            return min(6, max(1, int(m.group(1))))
    try:
        if style.builtin and re.fullmatch(r'\s*\d\s*', sid):
            return min(6, max(1, int(sid.strip())))
    except Exception:
        pass
    return None


def parse_docx(file_path, store, options):
    level = int(options.get('heading_level') or 2)
    meta = {'title': '', 'author': ''}
    try:
        import docx as pydocx
        d = pydocx.Document(file_path)
        cp = d.core_properties
        meta['title'] = (cp.title or '').strip()
        meta['author'] = (cp.author or '').strip()
    except Exception:
        pass

    html = None
    try:
        import mammoth

        def handle_image(image):
            with image.open() as image_file:
                data = image_file.read()
            ct = getattr(image, 'content_type', None) or 'image/png'
            ext = ct.split('/')[-1].split(';')[0]
            return {'src': store.add(data, ext)}

        with open(file_path, 'rb') as f:
            result = mammoth.convert_to_html(
                f, style_map=_style_map(),
                convert_image=mammoth.images.img_element(handle_image))
        html = result.value
    except Exception:
        html = None

    if not html or not html.strip():
        html = _fallback_html(file_path)

    from services.parser_html import split_html
    sections = split_html(html, level)
    if len(sections) <= 1:
        body = html.strip()
        if body:
            sections = [{'title': '全文', 'html': body}]
            return meta, sections, []
    toc = [{'title': s['title'], 'section': i, 'children': []} for i, s in enumerate(sections)]
    return meta, sections, toc


def _fallback_html(file_path):
    import docx as pydocx

    d = pydocx.Document(file_path)
    parts = []
    for para in d.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        h = _heading_level_of(para.style)
        if h and h <= 6:
            parts.append('<h%d>%s</h%d>' % (h, escape(text), h))
        else:
            parts.append('<p>%s</p>' % escape(text))
    for table in d.tables:
        rows = []
        for row in table.rows:
            cells = ['<td>%s</td>' % escape(c.text.strip()) for c in row.cells]
            rows.append('<tr>%s</tr>' % ''.join(cells))
        parts.append('<table>%s</table>' % ''.join(rows))
    return ''.join(parts)
