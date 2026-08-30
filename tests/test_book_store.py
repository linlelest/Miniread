import base64
import io
import os
import shutil
import unittest
import zipfile

from services import book_store
from services.book_store import (
    asset_path,
    build_canonical,
    load_manifest,
    needs_build,
    read_section,
    remove_canonical,
)
from services.sanitizer import sanitize_fragment

PNG_1PX = base64.b64decode(
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=='
)

_TMP = os.path.join(os.path.dirname(__file__), '_tmp')


class BookStoreTestBase(unittest.TestCase):
    def setUp(self):
        os.makedirs(_TMP, exist_ok=True)
        self._ids = []

    def tearDown(self):
        for bid in self._ids:
            remove_canonical(bid)
        shutil.rmtree(_TMP, ignore_errors=True)

    def _tmp_file(self, name, data):
        path = os.path.join(_TMP, name)
        mode = 'wb' if isinstance(data, bytes) else 'w'
        with open(path, mode, **({} if isinstance(data, bytes) else {'encoding': 'utf-8'})) as f:
            f.write(data)
        return path

    def _new_id(self):
        self._ids.append(88000 + len(self._ids) + int(os.urandom(2).hex(), 16) % 500)
        return self._ids[-1]


class TxtTests(BookStoreTestBase):
    def test_encodings_auto(self):
        text = '第一章 起点\n\n段落一。\n\n第二章 转折\n\n段落二。\n'
        for enc in ('utf-8', 'gb18030', 'utf-16'):
            bid = self._new_id()
            p = self._tmp_file('t_%s.txt' % enc, text.encode(enc))
            m = build_canonical(bid, 'txt', p, {'encoding': 'auto'})
            self.assertEqual(len(m['sections']), 2, enc)
            self.assertIn('第一章', m['sections'][0]['title'], enc)
            self.assertIn('段落一。', read_section(bid, 0), enc)
            self.assertIn('段落二。', read_section(bid, 1), enc)

    def test_custom_regex(self):
        text = '卷一 上\n内容A\n卷一 下\n内容B\n'
        bid = self._new_id()
        p = self._tmp_file('t.txt', text.encode('utf-8'))
        m = build_canonical(bid, 'txt', p, {'chapter_regex': r'^卷[一二]+ .+$'})
        self.assertEqual(len(m['sections']), 2)
        self.assertIn('内容A', read_section(bid, 0))
        self.assertIn('内容B', read_section(bid, 1))
        self.assertEqual(m['import_options'].get('chapter_regex'), r'^卷[一二]+ .+$')

    def test_invalid_custom_regex(self):
        bid = self._new_id()
        p = self._tmp_file('t.txt', 'x'.encode('utf-8'))
        with self.assertRaises(ValueError):
            build_canonical(bid, 'txt', p, {'chapter_regex': '('})

    def test_strip_toc(self):
        heads = ['第一章 甲', '第二章 乙', '第三章 丙', '第四章 丁', '第五章 戊']
        lines = list(heads) + ['']
        for i, h in enumerate(heads):
            lines.append(h)
            lines.append('正文第%d段。' % i)
        bid = self._new_id()
        p = self._tmp_file('t.txt', '\n'.join(lines).encode('utf-8'))
        m = build_canonical(bid, 'txt', p, {'strip_toc': True})
        self.assertEqual(len(m['sections']), 5)
        self.assertIn('正文第0段。', read_section(bid, 0))

    def test_hard_line_breaks(self):
        text = '第一章 起点\n\n段落A第一行。\n  缩进行。\n\n段落B。\n'
        bid = self._new_id()
        p = self._tmp_file('br.txt', text.encode('utf-8'))
        m = build_canonical(bid, 'txt', p, {})
        self.assertEqual(len(m['sections']), 1)
        self.assertEqual(m['toc'], [])
        html = read_section(bid, 0)
        self.assertIn('<div class="ln">段落A第一行。</div>', html)
        self.assertIn('<div class="ln">\u00a0\u00a0缩进行。</div>', html)
        self.assertIn('<div class="ln-gap"></div>', html)
        self.assertIn('<div class="ln">段落B。</div>', html)
        self.assertNotIn('<br/>', html)

    def test_no_match_whole_book(self):
        bid = self._new_id()
        p = self._tmp_file('t.txt', '没有章节标记\n第二行\n'.encode('utf-8'))
        m = build_canonical(bid, 'txt', p, {})
        self.assertEqual(len(m['sections']), 1)
        self.assertEqual(m['toc'], [])

    def test_english_chapters(self):
        text = 'Chapter 1 The Beginning\n\nText one.\n\nChapter 2 The End\n\nText two.\n'
        bid = self._new_id()
        p = self._tmp_file('en1.txt', text.encode('utf-8'))
        m = build_canonical(bid, 'txt', p, {})
        self.assertEqual(len(m['sections']), 2)

    def test_english_word_chapters(self):
        text = 'Chapter One\n\nA.\n\nChapter Twenty\n\nB.\n\nPrologue\n\nP.\n\nEpilogue\n\nE.\n'
        bid = self._new_id()
        p = self._tmp_file('en2.txt', text.encode('utf-8'))
        m = build_canonical(bid, 'txt', p, {})
        self.assertGreaterEqual(len(m['sections']), 3)

    def test_part_book_volume(self):
        text = 'Part One\n\nA.\n\nPart Two\n\nB.\n\nVolume 2\n\nC.\n'
        bid = self._new_id()
        p = self._tmp_file('en3.txt', text.encode('utf-8'))
        m = build_canonical(bid, 'txt', p, {})
        self.assertEqual(len(m['sections']), 3)

    def test_chinese_numbered_sections(self):
        text = '一、总览\n内容甲。\n二、详情\n内容乙。\n（三）补充\n内容丙。\n'
        bid = self._new_id()
        p = self._tmp_file('cn1.txt', text.encode('utf-8'))
        m = build_canonical(bid, 'txt', p, {})
        self.assertEqual(len(m['sections']), 3)


def _make_epub(path):
    buf = io.BytesIO()
    zi = zipfile.ZipInfo('mimetype')
    zi.compress_type = zipfile.ZIP_STORED
    with zipfile.ZipFile(buf, 'w') as z:
        z.writestr(zi, 'application/epub+zip')
        z.writestr('META-INF/container.xml', (
            '<?xml version="1.0"?><container version="1.0" '
            'xmlns="urn:oasis:names:tc:opendocument:xmlns:container"><rootfiles>'
            '<rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>'
            '</rootfiles></container>'))
        z.writestr('OEBPS/content.opf', (
            '<?xml version="1.0"?><package xmlns="http://www.idpf.org/2007/opf" '
            'version="2.0" unique-identifier="uid"><metadata xmlns:dc="http://purl.org/dc/elements/1.1/">'
            '<dc:title>测试书</dc:title><dc:creator>作者甲</dc:creator><dc:language>zh</dc:language>'
            '<meta name="cover" content="cover-img"/></metadata><manifest>'
            '<item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>'
            '<item id="c1" href="chap1.xhtml" media-type="application/xhtml+xml"/>'
            '<item id="c2" href="chap2.xhtml" media-type="application/xhtml+xml"/>'
            '<item id="cover-img" href="images/cover.png" media-type="image/png"/>'
            '<item id="pic" href="images/pic.png" media-type="image/png"/>'
            '</manifest><spine toc="ncx"><itemref idref="c1"/><itemref idref="c2"/></spine></package>'))
        z.writestr('OEBPS/chap1.xhtml', (
            '<html><head><title>c1</title></head><body><h1>第一章 起点</h1>'
            '<p>这是第一章的内容。</p><img src="images/pic.png"/></body></html>'))
        z.writestr('OEBPS/chap2.xhtml', (
            '<html><head><title>c2</title></head><body><h1>第二章 转折</h1>'
            '<p>这是第二章的内容。</p></body></html>'))
        z.writestr('OEBPS/toc.ncx', (
            '<?xml version="1.0"?><ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">'
            '<head/><docTitle><text>测试书</text></docTitle><navMap>'
            '<navPoint id="n1" playOrder="1"><navLabel><text>第一章 起点</text></navLabel>'
            '<content src="chap1.xhtml"/></navPoint>'
            '<navPoint id="n2" playOrder="2"><navLabel><text>第二章 转折</text></navLabel>'
            '<content src="chap2.xhtml"/></navPoint></navMap></ncx>'))
        z.writestr('OEBPS/images/cover.png', PNG_1PX)
        z.writestr('OEBPS/images/pic.png', PNG_1PX)
    with open(path, 'wb') as f:
        f.write(buf.getvalue())


class EpubTests(BookStoreTestBase):
    def test_epub_full(self):
        bid = self._new_id()
        p = self._tmp_file('t.epub', b'')
        _make_epub(p)
        m = build_canonical(bid, 'epub', p, {})
        self.assertEqual(m['title'], '测试书')
        self.assertEqual(m['author'], '作者甲')
        self.assertEqual(len(m['sections']), 2)
        self.assertIn('这是第一章的内容。', read_section(bid, 0))
        self.assertIn('/asset/', read_section(bid, 0))
        assets = os.listdir(os.path.join(book_store.book_dir(bid), 'assets'))
        self.assertEqual(len(assets), 2)
        self.assertIsNotNone(m['cover'])
        self.assertGreaterEqual(len(m['toc']), 2)
        self.assertEqual(m['toc'][0]['section'], 0)
        self.assertEqual(m['toc'][1]['section'], 1)


class DocxTests(BookStoreTestBase):
    def test_docx_real_chapters(self):
        import docx as pydocx
        d = pydocx.Document()
        d.add_heading('第一章 甲', level=1)
        d.add_paragraph('甲内容独有。')
        d.add_heading('第二章 乙', level=1)
        d.add_paragraph('乙内容独有。')
        d.add_heading('第三章 丙', level=1)
        d.add_paragraph('丙内容独有。')
        bid = self._new_id()
        p = os.path.join(_TMP, 't.docx')
        d.save(p)
        m = build_canonical(bid, 'docx', p, {'heading_level': 2})
        self.assertEqual(len(m['sections']), 3)
        s0 = read_section(bid, 0)
        self.assertIn('甲内容独有。', s0)
        self.assertNotIn('乙内容独有。', s0)
        self.assertNotIn('丙内容独有。', s0)

    def test_docx_no_headings_single_section(self):
        import docx as pydocx
        d = pydocx.Document()
        d.add_paragraph('第一段，没有标题。')
        d.add_paragraph('第二段。')
        bid = self._new_id()
        p = os.path.join(_TMP, 'plain.docx')
        d.save(p)
        m = build_canonical(bid, 'docx', p, {})
        self.assertEqual(len(m['sections']), 1)
        self.assertEqual(m['toc'], [])
        self.assertIn('第一段', read_section(bid, 0))
        self.assertIn('第二段', read_section(bid, 0))


def _make_fb2(path):
    b64 = base64.b64encode(PNG_1PX).decode()
    xml = (
        '<?xml version="1.0" encoding="utf-8"?>'
        '<FictionBook xmlns="http://www.gribuser.ru/xml/fictionbook/2.0" '
        'xmlns:l="http://www.w3.org/1999/xlink">'
        '<description><title-info><book-title>FB2书</book-title>'
        '<author><first-name>名</first-name><last-name>姓</last-name></author>'
        '<lang>zh</lang><coverpage><image l:href="#cover.png"/></coverpage></title-info></description>'
        '<body>'
        '<section><title><p>第一章 起</p></title><p>FB2第一段。</p><image l:href="#pic.png"/></section>'
        '<section><title><p>第二章 承</p></title><p>FB2第二段。</p></section>'
        '</body>'
        '<binary id="cover.png" content-type="image/png">%s</binary>'
        '<binary id="pic.png" content-type="image/png">%s</binary>'
        '</FictionBook>' % (b64, b64)
    )
    with open(path, 'w', encoding='utf-8') as f:
        f.write(xml)


class Fb2Tests(BookStoreTestBase):
    def test_fb2_sections_images_cover(self):
        bid = self._new_id()
        p = self._tmp_file('t.fb2', '')
        _make_fb2(p)
        m = build_canonical(bid, 'fb2', p, {})
        self.assertEqual(m['title'], 'FB2书')
        self.assertEqual(m['author'], '名 姓')
        self.assertEqual(len(m['sections']), 2)
        s0 = read_section(bid, 0)
        self.assertIn('FB2第一段。', s0)
        self.assertIn('/asset/', s0)
        self.assertIsNotNone(m['cover'])


class HtmlMdRtfTests(BookStoreTestBase):
    def test_html_nested_headings(self):
        bid = self._new_id()
        html = ('<html><head><title>HTML书</title></head><body>'
                '<div><h2>甲节</h2><p>内容甲</p></div>'
                '<div><h2>乙节</h2><p>内容乙</p></div></body></html>')
        p = self._tmp_file('t.html', html)
        m = build_canonical(bid, 'html', p, {'heading_level': 2})
        self.assertEqual(m['title'], 'HTML书')
        self.assertEqual(len(m['sections']), 2)
        self.assertIn('内容甲', read_section(bid, 0))

    def test_md_headings(self):
        bid = self._new_id()
        md = '## 第一节\n\n内容A\n\n## 第二节\n\n内容B\n'
        p = self._tmp_file('t.md', md)
        m = build_canonical(bid, 'md', p, {'heading_level': 2})
        self.assertEqual(len(m['sections']), 2)
        self.assertIn('内容A', read_section(bid, 0))
        self.assertIn('内容B', read_section(bid, 1))

    def test_rtf(self):
        bid = self._new_id()
        rtf = r'{\rtf1\ansi 第一章 起点\par 段落一。\par 第二章 转折\par 段落二。\par}'
        p = self._tmp_file('t.rtf', rtf)
        m = build_canonical(bid, 'rtf', p, {})
        self.assertGreaterEqual(len(m['sections']), 1)


class SanitizerTests(unittest.TestCase):
    def test_removes_dangerous(self):
        html = ('<p onclick="evil()">a</p>'
                '<script>bad()</script>'
                '<a href="javascript:alert(1)">x</a>'
                '<a href="https://ext.example">外链</a>'
                '<iframe src="http://x"></iframe>'
                '<!-- comment -->')
        out = sanitize_fragment(html)
        self.assertNotIn('<script', out)
        self.assertNotIn('iframe', out)
        self.assertNotIn('onclick', out)
        self.assertNotIn('javascript:', out)
        self.assertIn('href="#"', out)
        self.assertIn('data-href="https://ext.example"', out)
        self.assertNotIn('<!--', out)


class AssetAndStateTests(BookStoreTestBase):
    def test_asset_path_traversal(self):
        bid = self._new_id()
        store = book_store.AssetStore(bid)
        store.add(PNG_1PX, 'png')
        ok = asset_path(bid, os.listdir(os.path.join(book_store.book_dir(bid), 'assets'))[0])
        self.assertIsNotNone(ok)
        self.assertIsNone(asset_path(bid, '../secret.txt'))
        self.assertIsNone(asset_path(bid, 'sub/dir.png'))
        self.assertIsNone(asset_path(bid, '.hidden'))
        self.assertIsNone(asset_path(bid, ''))
        self.assertIsNone(asset_path(bid, 'missing.png'))

    def test_needs_build(self):
        bid = self._new_id()
        p = self._tmp_file('t.txt', '第一章 a\n正文\n'.encode('utf-8'))
        build_canonical(bid, 'txt', p, {})
        self.assertFalse(needs_build({'canonical_status': 'ready', 'id': bid}))
        self.assertTrue(needs_build({'canonical_status': 'legacy', 'id': bid}))
        self.assertTrue(needs_build({'canonical_status': 'pending', 'id': bid}))
        self.assertTrue(needs_build({'canonical_status': None, 'id': bid}))
        remove_canonical(bid)
        self.assertTrue(needs_build({'canonical_status': 'ready', 'id': bid}))

    def test_unsupported_format_rejected(self):
        bid = self._new_id()
        p = self._tmp_file('t.lit', b'MZBAD')
        with self.assertRaises(ValueError):
            build_canonical(bid, 'lit', p, {})


if __name__ == '__main__':
    unittest.main()
